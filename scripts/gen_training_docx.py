#!/usr/bin/env python3
# Tạo file presentation DOCX nội dung đào tạo (4 module)
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

NAVY = RGBColor(0x09, 0x37, 0x74)
TEAL = RGBColor(0x15, 0xB5, 0xB0)
CORAL = RGBColor(0xFF, 0x68, 0x59)
DARK = RGBColor(0x2C, 0x33, 0x5D)
GRAY = RGBColor(0x94, 0xA3, 0xB8)

modules = [
    {
        "title": "Module 1 · Định hướng & giới thiệu chương trình",
        "summary": "Module mở đầu giúp cả hai bên nắm rõ cách chương trình vận hành, cam kết cần có, và vai trò của Điều phối viên (ĐPV).",
        "slides": [
            ("Chương trình mentoring cộng đồng là gì?", [
                "Lộ trình đồng hành 9 tháng có cấu trúc giữa mentor (người dẫn dắt) và mentee (người được đồng hành).",
                "Không phải cố vấn rời rạc — mà là một hành trình có mục tiêu, có điểm chạm, có đo lường.",
                "Được hỗ trợ xuyên suốt bởi Điều phối viên (ĐPV).",
            ]),
            ("Vai trò & kỳ vọng mỗi bên", [
                "Mentor: chia sẻ kinh nghiệm, đặt câu hỏi dẫn dắt, không làm thay việc của mentee.",
                "Mentee: chủ động đặt mục tiêu, chuẩn bị trước buổi gặp, thực hành và phản hồi.",
                "ĐPV: theo dõi tiến độ, gỡ khó khi có tạm dừng hoặc vướng mắc.",
            ]),
            ("3 giá trị cốt lõi", [
                "Tôn trọng & Lắng nghe thật — không phán xét, đặt mình vào vị trí người kia.",
                "Cam kết & Trách nhiệm — giữ lịch hẹn, đồng hành trọn vẹn, không bỏ ngang.",
                "Cởi mở & Trung thực — chia sẻ thật, sẵn sàng nhận góp ý.",
            ]),
            ("Lộ trình 9 tháng", [
                "Tháng 1: Kết nối lần đầu + ký thoả thuận đồng hành.",
                "Tháng 2–8: Gặp định kỳ, ghi chú, phản tư hằng tháng.",
                "Tháng 9: Tổng kết hành trình, phản tư kết thúc.",
            ]),
        ],
    },
    {
        "title": "Module 2 · Kỹ năng lắng nghe & đặt câu hỏi",
        "summary": "Kỹ năng cốt lõi giúp mentor dẫn dắt thay vì áp đặt, và mentee mở lòng chia sẻ đúng trọng tâm.",
        "slides": [
            ("Tại sao lắng nghe quan trọng?", [
                "Lắng nghe thật giúp hiểu đúng vấn đề, không chỉ nghe để đáp.",
                "Người được lắng nghe cảm thấy được tôn trọng và an toàn để chia sẻ.",
                "Phần lớn giải pháp đến từ chính người nói khi được hỏi đúng cách.",
            ]),
            ("Lắng nghe chủ động (Active Listening)", [
                "Duy trì giao tiếp mắt, gật đầu, tóm tắt lại điều vừa nghe.",
                "Dùng câu phản hồi: \"Vậy điều bạn đang lo là...?\"",
                "Không ngắt lời, không vội đưa giải pháp khi chưa hiểu trọn vẹn.",
            ]),
            ("Đặt câu hỏi mở", [
                "Tránh câu hỏi đóng (chỉ trả lời có/không).",
                "Ưu tiên: \"Điều gì khiến bạn nghĩ vậy?\", \"Bạn đã thử những gì?\", \"Kết quả bạn mong muốn là gì?\"",
                "Dùng câu hỏi để giúp mentee tự tìm ra hướng đi.",
            ]),
            ("Những điều cần tránh", [
                "Phán xét, so sánh, áp đặt kinh nghiệm cá nhân quá sớm.",
                "Chuyển chủ đề sang câu chuyện của mình quá nhanh.",
                "Đưa ra lời khuyên khi chưa được hỏi.",
            ]),
        ],
    },
    {
        "title": "Module 3 · Thiết lập mục tiêu & thoả thuận đồng hành",
        "summary": "Mục tiêu rõ ràng + thoả thuận minh bạch là nền tảng để 9 tháng đi đúng hướng và đo lường được.",
        "slides": [
            ("Vì sao cần mục tiêu rõ ràng?", [
                "Mục tiêu mơ hồ dẫn đến buổi gặp lạc hướng, mất động lực.",
                "Mục tiêu rõ giúp cả hai biết mình đang làm gì và vì sao.",
                "Là thước đo để đánh giá hành trình có hiệu quả không.",
            ]),
            ("Mô hình SMART", [
                "Specific — cụ thể, rõ ràng.",
                "Measurable — đo lường được.",
                "Achievable — khả thi.",
                "Relevant — liên quan đến mục tiêu lớn của mentee.",
                "Time-bound — có mốc thời gian.",
            ]),
            ("Xây dựng thoả thuận đồng hành", [
                "Thống nhất tần suất gặp, hình thức (offline/online), thời lượng.",
                "Quy tắc giao tiếp, bảo mật thông tin.",
                "Cơ chế khi có vấn đề: liên hệ ĐPV, tạm dừng có lộ trình.",
            ]),
            ("Cam kết hai chiều", [
                "Mentor cam kết thời gian, sự chuẩn bị và tính nhất quán.",
                "Mentee cam kết chủ động, thực hành và phản hồi đúng hạn.",
                "Ghi lại thoả thuận để cả hai cùng nhìn lại khi cần.",
            ]),
        ],
    },
    {
        "title": "Module 4 · Phản tư & phát triển liên tục",
        "summary": "Phản tư là trái tim của chương trình — giúp cặp đồng hành nhìn lại, điều chỉnh và trưởng thành hơn mỗi tháng.",
        "slides": [
            ("Phản tư là gì?", [
                "Khoảng dừng có chủ đích mỗi tháng để nhìn lại hành trình.",
                "Không phải đánh giá đúng/sai — mà là nhận diện điều đang diễn ra.",
                "Gồm 4 trạng thái cảm nhận: kết nối tốt, tìm nhịp, chưa ổn, cần hỗ trợ.",
            ]),
            ("Cách thực hiện phản tư hiệu quả", [
                "Trả lời trung thực cảm nhận tháng vừa qua.",
                "Ghi rõ điều đã làm được, điều còn vướng.",
                "Đặt 1 mục tiêu nhỏ cho tháng tiếp theo.",
            ]),
            ("Khi phản tư báo hiệu màu vàng/đỏ", [
                "Màu vàng (chưa ổn) 2 tháng liên tiếp → ĐPV chủ động liên hệ.",
                "Màu đỏ (cần hỗ trợ) → tạo yêu cầu hỗ trợ ngay.",
                "Mục tiêu: can thiệp sớm, không để vấn đề lớn dần.",
            ]),
            ("Vòng lặp phát triển", [
                "Hành động → Phản tư → Nhận ra → Điều chỉnh → Hành động tốt hơn.",
                "Mỗi tháng là một vòng lặp nhỏ tiến dần đến mục tiêu 9 tháng.",
                "Ghi lại vào nhật ký hành trình để thấy được sự trưởng thành.",
            ]),
        ],
    },
]

doc = Document()

# Trang bìa
title = doc.add_heading("Chương trình Đào tạo Mentor", level=0)
for run in title.runs:
    run.font.color.rgb = NAVY
sub = doc.add_paragraph("Phác thảo nội dung 4 module · Nền Tảng Mentoring Cộng Đồng")
sub.runs[0].font.color.rgb = GRAY
sub.runs[0].font.size = Pt(12)

doc.add_paragraph()

for mod in modules:
    doc.add_page_break()
    # Tiêu đề module
    h = doc.add_heading(mod["title"], level=1)
    for run in h.runs:
        run.font.color.rgb = NAVY
    # Summary
    p = doc.add_paragraph(mod["summary"])
    p.runs[0].font.italic = True
    p.runs[0].font.color.rgb = DARK

    for i, (slide_title, bullets) in enumerate(mod["slides"], 1):
        doc.add_paragraph()
        sh = doc.add_heading(f"{i}. {slide_title}", level=2)
        for run in sh.runs:
            run.font.color.rgb = TEAL
        for b in bullets:
            bp = doc.add_paragraph(b, style="List Bullet")
            bp.runs[0].font.color.rgb = DARK

out = "/root/.openclaw/workspace/mentor-platform/Mentor_Training_Content.docx"
doc.save(out)
print("Đã tạo:", out)
